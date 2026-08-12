"""
בדיקה מינימלית: האם python-docx יוצר RTL נכון על המחשב הזה?
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

doc = Document()

# ── פסקה 1: RTL ידני ──
para = doc.add_paragraph()
pPr = para._p.get_or_add_pPr()

bidi = OxmlElement('w:bidi')
pPr.append(bidi)

jc = OxmlElement('w:jc')
jc.set(qn('w:val'), 'right')
pPr.append(jc)

run = para.add_run('שלום — זה צריך להיות מיושר לימין!')
run.font.size = Pt(14)

rPr = run._r.get_or_add_rPr()
rtl_el = OxmlElement('w:rtl')
rPr.append(rtl_el)
lang = OxmlElement('w:lang')
lang.set(qn('w:bidi'), 'he-IL')
rPr.append(lang)

# ── פסקה 2: LTR רגיל ──
para2 = doc.add_paragraph()
run2 = para2.add_run('This should be left-aligned (LTR)')
run2.font.size = Pt(14)

# 1. RTL ב-Section
from docx.oxml import OxmlElement as OE
sectPr = doc.sections[0]._sectPr
sectPr.append(OE('w:bidi'))

# 2. RTL בהגדרות המסמך עצמו (document settings)
settings_el = doc.settings.element
settings_el.append(OE('w:bidi'))

# 3. הסרת Compatibility Mode — גורם ל-Word לפתוח כ-Modern
from docx.oxml.ns import qn as _qn
compat = settings_el.find(_qn('w:compat'))
if compat is not None:
    settings_el.remove(compat)

doc.save('test_rtl2.docx')
print("נשמר: test_rtl2.docx")

# הדפסת ה-XML שנוצר לפסקה הראשונה
import zipfile, re
with zipfile.ZipFile('test_rtl2.docx', 'r') as z:
    xml = z.read('word/document.xml').decode('utf-8')

# מציאת ה-pPr הראשון
match = re.search(r'<w:pPr>.*?</w:pPr>', xml, re.DOTALL)
if match:
    print("\nXML של הפסקה הראשונה:")
    print(match.group(0))
else:
    print("\nלא נמצא pPr — הפסקה ריקה מ-XML")
