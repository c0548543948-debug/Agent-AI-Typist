"""
document_writer.py — יצירת קובץ Word מתוצאות התמלול

האחריות של הקובץ הזה:
1. לקבל רשימת TranscriptionResult (תמלול של כל העמודים)
2. לבנות מסמך Word עם כיווניות נכונה (RTL לעברית, LTR לאנגלית)
3. להדגיש מילים לא-ודאיות [מילה?] בצבע צהוב
4. להוסיף מפריד בין עמודים
5. להוסיף טבלת סיכום בסוף המסמך
6. לשמור לקובץ .docx
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from loguru import logger

from config import AppConfig
from ocr_engine import TranscriptionResult


# ============================================================
# פונקציה ראשית
# ============================================================

def create_docx(
    results: list[TranscriptionResult],
    output_path: str,
    config: AppConfig,
) -> None:
    """
    מקבלת רשימת תוצאות תמלול ויוצרת קובץ Word.

    results     — רשימה מסודרת של TranscriptionResult (עמוד אחד כל אחד)
    output_path — נתיב לשמירת קובץ .docx
    config      — הגדרות המערכת (גופן, גודל, צבעים)
    """
    doc = Document()

    # הגדרת RTL ברמת המסמך כולו — זה מה שגורם ל-Word לכבד RTL
    _set_document_rtl_defaults(doc)

    # הגדרת שוליים קטנים יותר — ברירת המחדל של Word רחבה מדי
    _set_margins(doc, top=1.0, bottom=1.0, left=1.2, right=1.2)

    total_uncertain = 0

    for i, result in enumerate(results):
        # הוספת מפריד בין עמודים (לא לפני הראשון)
        if i > 0:
            _add_page_separator(doc, result.page_number)

        # כתיבת הטקסט של העמוד
        uncertain_count = _write_page_text(doc, result, config)
        total_uncertain += uncertain_count

        logger.info(
            f"עמוד {result.page_number}: נכתב "
            f"({uncertain_count} מילים לא-ודאיות)"
        )

    # טבלת סיכום בסוף המסמך
    _add_summary_table(doc, results, total_uncertain, config)

    # שמירה לדיסק
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    logger.info(f"מסמך נשמר: {output}")


# ============================================================
# כתיבת טקסט עמוד
# ============================================================

def _write_page_text(
    doc: Document,
    result: TranscriptionResult,
    config: AppConfig,
) -> int:
    """
    כותבת את הטקסט של עמוד אחד למסמך.
    מחלקת לפסקאות לפי שורות ריקות, ולכל פסקה קובעת כיווניות.
    מחזירה כמות מילים לא-ודאיות שנמצאו.
    """
    uncertain_total = 0

    # פיצול לפסקאות לפי שורות ריקות
    # שורה ריקה בין פסקאות היא הקונבנציה שביקשנו מ-Gemini
    paragraphs = re.split(r'\n\s*\n', result.text)

    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        # קביעת כיווניות הפסקה
        # אם מצב השפה הוא auto — מזהים לפי תוכן הפסקה עצמה
        # אם השתמשנו במצב he/en — עוקבים אחריו
        direction = _detect_paragraph_direction(para_text, result.detected_language)

        # כתיבת הפסקה עם הדגשת מילים לא-ודאיות
        uncertain_count = _add_paragraph(doc, para_text, direction, config)
        uncertain_total += uncertain_count

    return uncertain_total


def _detect_paragraph_direction(text: str, doc_language: str) -> str:
    """
    קובעת האם פסקה היא RTL (עברית) או LTR (אנגלית).

    אם שפת המסמך מוגדרת במפורש — מכבדים אותה.
    אחרת — בודקים נוכחות תווים עבריים בפסקה עצמה.
    """
    if doc_language == "he":
        return "rtl"
    if doc_language == "en":
        return "ltr"

    # מצב auto: פסקה עם תווים עבריים → RTL
    has_hebrew = bool(re.search(r'[א-תיִ-פֿ]', text))
    return "rtl" if has_hebrew else "ltr"


# ============================================================
# הוספת פסקה עם הדגשת מילים לא-ודאיות
# ============================================================

UNCERTAIN_PATTERN = re.compile(r'(\[[^\]]+\?\])')
"""
תבנית regex לזיהוי מילים לא-ודאיות.
מחפשת: [כלשהו?] — סוגריים מרובעות עם סימן שאלה בסוף.
"""

LEADING_PUNCT = re.compile(r'^([!?.،؟]+)(.*)')
"""תבנית לסימני פיסוק בתחילת שורה עברית — צריך להזיז אותם לסוף."""


def _fix_hebrew_punctuation(text: str) -> str:
    """
    מתקנת מיקום סימני פיסוק בפסקה עברית.

    Gemini מתמלל "!היי שלום" עם "!" בהתחלה של המחרוזת הלוגית.
    בפסקה ממויינת ימין (ללא RTL אמיתי), "!" מגיע לפני הטקסט — שגיאה ויזואלית.
    הפתרון: "!היי שלום" → "היי שלום!" כך שה-! יופיע בסוף.
    """
    fixed_lines = []
    for line in text.split('\n'):
        m = LEADING_PUNCT.match(line.strip())
        if m:
            punct, rest = m.group(1), m.group(2).strip()
            line = rest + punct if rest else punct
        fixed_lines.append(line)
    return '\n'.join(fixed_lines)


def _add_paragraph(
    doc: Document,
    text: str,
    direction: str,
    config: AppConfig,
) -> int:
    """
    מוסיפה פסקה למסמך.

    מחלקת את הטקסט לחלקים: חלקים רגילים וחלקים לא-ודאיים [מילה?].
    לכל חלק לא-ודאי מוסיפה הדגשת רקע צהובה.
    מחזירה כמות מילים לא-ודאיות בפסקה.
    """
    # תיקון מיקום סימני פיסוק בעברית
    if direction == "rtl":
        text = _fix_hebrew_punctuation(text)

    para = doc.add_paragraph()
    _set_paragraph_direction(para, direction)

    # פיצול הטקסט לחלקים לפי תבנית [מילה?]
    # דוגמה: "שלום [עולם?] טוב" → ["שלום ", "[עולם?]", " טוב"]
    parts = UNCERTAIN_PATTERN.split(text)
    uncertain_count = 0

    for part in parts:
        if not part:
            continue

        run = para.add_run(part)
        run.font.name = config.default_font
        run.font.size = Pt(config.default_font_size / 2)
        # הערה: default_font_size ב-config הוא ב-half-points (24 = 12pt)
        # Pt() מקבל נקודות רגילות, אז מחלקים ב-2

        # בדיקה אם החלק הוא מילה לא-ודאית
        if UNCERTAIN_PATTERN.match(part):
            _highlight_yellow(run)
            uncertain_count += 1

    return uncertain_count


# ============================================================
# עיצוב XML של Word
# ============================================================

def _set_paragraph_direction(para, direction: str) -> None:
    """
    מגדירה יישור לפסקה.

    במקום XML מורכב של RTL/bidi שלא עובד באופן עקבי ב-Word,
    משתמשים ב-API הפשוט של python-docx:
    - עברית → ישור ימין (RIGHT)
    - אנגלית → ישור שמאל (LEFT)

    Word מזהה עברית Unicode אוטומטית ומטפל בכיוון האותיות בתוך השורה.
    """
    if direction == "rtl":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT



def _highlight_yellow(run) -> None:
    """
    מוסיפה הדגשת רקע צהובה ל-run (קטע טקסט).

    למה XML ישירות?
    אותה סיבה כמו כיווניות — python-docx לא חושף highlight API פשוט.
    Word מאחסן הדגשה בתגית <w:highlight w:val="yellow"/> בתוך <w:rPr>.
    """
    rPr = run._r.get_or_add_rPr()
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), 'yellow')
    rPr.append(highlight)


# ============================================================
# מפריד בין עמודים
# ============================================================

def _add_page_separator(doc: Document, page_number: int) -> None:
    """
    מוסיפה מפריד ויזואלי בין עמודים:
    שורה ריקה + "─── עמוד X ───" מרוכז + שורה ריקה.

    למה לא page break?
    Page break ב-Word מתחיל עמוד פיזי חדש.
    אנחנו רוצים שהכל יהיה באותו מסמך רציף אך עם סימון ברור.
    """
    doc.add_paragraph()  # שורה ריקה לפני

    sep = doc.add_paragraph(f"─── עמוד {page_number} ───")
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.runs[0]
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)  # אפור

    doc.add_paragraph()  # שורה ריקה אחרי


# ============================================================
# טבלת סיכום
# ============================================================

def _add_summary_table(
    doc: Document,
    results: list[TranscriptionResult],
    total_uncertain: int,
    config: AppConfig,
) -> None:
    """
    מוסיפה טבלת סיכום בסוף המסמך.
    מכילה: עמודים שעובדו, סה"כ מילים לא-ודאיות, זמן עיבוד כולל.
    """
    doc.add_paragraph()  # רווח לפני הטבלה

    heading = doc.add_paragraph("סיכום עיבוד")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.bold = True
    heading.runs[0].font.size = Pt(10)
    _set_paragraph_direction(heading, "rtl")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # כותרת
    header_cells = table.rows[0].cells
    header_cells[0].text = "פריט"
    header_cells[1].text = "ערך"

    # נתונים
    total_ms = sum(r.api_call_ms for r in results)
    failed = sum(1 for r in results if not r.success)

    rows_data = [
        ("עמודים שעובדו", str(len(results))),
        ("עמודים שנכשלו", str(failed)),
        ("מילים לא-ודאיות", str(total_uncertain)),
        ("זמן עיבוד כולל", f"{total_ms / 1000:.1f} שניות"),
    ]

    for label, value in rows_data:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value


# ============================================================
# שוליים
# ============================================================

def _set_document_rtl_defaults(doc: Document) -> None:
    """
    פונקציה ריקה — לא מגדירים RTL ברמת המסמך.
    Word מזהה עברית Unicode אוטומטית ומטפל בכיוון האותיות.
    ישור ימין מוגדר ברמת כל פסקה בנפרד.
    """
    pass


def _set_margins(doc: Document, top: float, bottom: float, left: float, right: float) -> None:
    """
    מגדירה שוליים למסמך (באינצ'ים).
    ברירת המחדל של Word היא 1 אינץ' מכל צד — לפעמים רחב מדי לטקסט ארוך.
    """
    section = doc.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
